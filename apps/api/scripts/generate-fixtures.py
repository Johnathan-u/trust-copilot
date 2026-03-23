"""Generate XLSX/DOCX fixture files from JSON definitions (EXP-09)."""

import json
import sys
from pathlib import Path

API_ROOT = Path(__file__).resolve().parent.parent
FIXTURES = API_ROOT / "tests" / "fixtures" / "questionnaires"


def generate_xlsx_from_json(json_path: Path, out_path: Path) -> None:
    """Create minimal XLSX from JSON. Places section headers before first question of each section."""
    import openpyxl

    data = json.loads(json_path.read_text())
    wb = openpyxl.Workbook()
    ws = wb.active
    sheet_name = data[0]["source_location"].get("sheet", "Sheet1") if data else "Sheet1"
    ws.title = sheet_name[:31]
    ws.cell(1, 1, "Question")
    ws.cell(1, 2, "Answer")
    prev_section = None
    for q in data:
        loc = q.get("source_location", {})
        r, c = loc.get("row", 1), loc.get("col", 1)
        section = q.get("section")
        if section and section != prev_section:
            ws.cell(r - 1, 1, section)
            prev_section = section
        ws.cell(r, c, q.get("text", ""))
        if c == 1:
            ws.cell(r, 2, "")
    wb.save(out_path)


def generate_docx_from_json(json_path: Path, out_path: Path) -> None:
    """Create minimal DOCX from JSON question definitions."""
    from docx import Document

    data = json.loads(json_path.read_text())
    doc = Document()
    section = None
    for q in data:
        if q.get("section") != section:
            section = q.get("section")
            if section:
                doc.add_paragraph(section)
        doc.add_paragraph(q["text"])
        doc.add_paragraph("")
    doc.save(out_path)


def main() -> None:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    for name in ["simple_soc2", "iso27001", "vendor_dd"]:
        j = FIXTURES / f"{name}.json"
        if not j.exists():
            continue
        generate_xlsx_from_json(j, FIXTURES / f"{name}.xlsx")
        generate_docx_from_json(j, FIXTURES / f"{name}.docx")
        print(f"Generated {name}.xlsx, {name}.docx")


if __name__ == "__main__":
    main()
    sys.exit(0)
