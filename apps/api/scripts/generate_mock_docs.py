"""Generate mock evidence documents for demo/testing (DOC-01)."""

from pathlib import Path

API_ROOT = Path(__file__).resolve().parent.parent
OUT = API_ROOT / "tests" / "fixtures" / "documents"


def generate_security_controls_docx(out_path: Path) -> None:
    """Create security controls overview DOCX."""
    from docx import Document
    doc = Document()
    doc.add_heading("Information Security Controls Overview", 0)
    doc.add_paragraph(
        "This document describes the security controls maintained by our organization."
    )
    doc.add_heading("1. Written Information Security Policy", level=1)
    doc.add_paragraph(
        "Yes. The company maintains a written information security policy that is "
        "approved by management, communicated to all personnel, and reviewed annually."
    )
    doc.add_heading("2. Designated Security Responsibility", level=1)
    doc.add_paragraph(
        "Yes. A designated individual (CISO) is responsible for the security program, "
        "including policy development, risk assessment, and oversight."
    )
    doc.add_heading("3. Risk Assessments", level=1)
    doc.add_paragraph(
        "Yes. The organization performs annual risk assessments to identify threats "
        "and vulnerabilities. Findings are documented and remediated according to priority."
    )
    doc.add_heading("4. Access Control", level=1)
    doc.add_paragraph(
        "Access to systems and data is controlled via role-based access control (RBAC). "
        "Principle of least privilege is applied. Access reviews are conducted quarterly."
    )
    doc.add_heading("5. Incident Response", level=1)
    doc.add_paragraph(
        "An incident response plan is maintained and tested annually. Security incidents "
        "are logged, triaged, and escalated per defined procedures."
    )
    doc.save(out_path)


def generate_sample_evidence_xlsx(out_path: Path) -> None:
    """Create sample evidence XLSX with policy/control content."""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Security Evidence"
    ws.cell(1, 1, "Control")
    ws.cell(1, 2, "Description")
    rows = [
        ("CC1.1", "Written information security policy is maintained and reviewed annually."),
        ("CC1.2", "Designated security officer (CISO) is responsible for the security program."),
        ("CC2.1", "Annual risk assessments are performed and documented."),
        ("CC2.2", "Risk register tracks identified risks and mitigation status."),
        ("CC3.1", "Access control policies limit access based on job function."),
        ("CC4.1", "Incident response plan is documented and tested annually."),
    ]
    for i, (ctrl, desc) in enumerate(rows, start=2):
        ws.cell(i, 1, ctrl)
        ws.cell(i, 2, desc)
    wb.save(out_path)


def generate_security_policy_pdf(out_path: Path) -> None:
    """Create minimal security policy PDF via pymupdf."""
    try:
        import pymupdf as fitz
    except ImportError:
        return
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    text = (
        "Information Security Policy\n\n"
        "1. Policy Statement: The company maintains a written information security "
        "policy approved by management.\n\n"
        "2. Security Officer: A designated individual (CISO) is responsible for "
        "the security program.\n\n"
        "3. Risk Assessment: Annual risk assessments are performed.\n\n"
        "4. Access Control: Role-based access control is applied. Access reviews "
        "occur quarterly.\n\n"
        "5. Incident Response: An incident response plan exists and is tested annually."
    )
    rect = fitz.Rect(50, 50, 545, 792)
    page.insert_textbox(rect, text, fontsize=11)
    doc.save(out_path)
    doc.close()


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    generate_security_controls_docx(OUT / "security_controls_overview.docx")
    generate_sample_evidence_xlsx(OUT / "sample_evidence.xlsx")
    generate_security_policy_pdf(OUT / "security_policy.pdf")
    print("Generated mock evidence documents:")
    for f in sorted(OUT.iterdir()):
        print(f"  {f.name}")


if __name__ == "__main__":
    main()
