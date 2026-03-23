"""TEST-05: Golden tests for export behavior."""

import json
from io import BytesIO
from pathlib import Path

import openpyxl
import pytest

from app.services.docx_writer import create_export_from_questionnaire_docx
from app.services.xlsx_writer import create_export_from_questionnaire, placements_from_questions_and_answers

FIXTURES = Path(__file__).resolve().parent / "fixtures" / "questionnaires"


def test_placements_from_questions_xlsx() -> None:
    """Placements builder produces correct sheet/row/col from source_location."""
    questions = [
        {"id": 1, "source_location": json.dumps({"sheet": "CC1", "row": 2, "col": 1})},
        {"id": 2, "source_location": json.dumps({"sheet": "CC1", "row": 3, "col": 1})},
    ]
    qa = {1: "Yes", 2: "No"}
    placements = placements_from_questions_and_answers(questions, qa)
    assert len(placements) == 2
    assert placements[0] == {"sheet": "CC1", "row": 2, "col": 1, "text": "Yes"}
    assert placements[1] == {"sheet": "CC1", "row": 3, "col": 1, "text": "No"}


def test_xlsx_export_in_memory() -> None:
    """Export creates workbook with answers in correct cells."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "CC1"
    ws.cell(2, 1, "Q1?")
    ws.cell(2, 2, "")
    ws.cell(3, 1, "Q2?")
    ws.cell(3, 2, "")
    buf = BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()

    questions = [
        {"id": 1, "source_location": json.dumps({"sheet": "CC1", "row": 2, "col": 1})},
        {"id": 2, "source_location": json.dumps({"sheet": "CC1", "row": 3, "col": 1})},
    ]
    qa = {1: "Answer 1", 2: "Answer 2"}
    out = create_export_from_questionnaire(xlsx_bytes, questions, qa)
    out_wb = openpyxl.load_workbook(BytesIO(out), data_only=False)
    assert out_wb["CC1"].cell(2, 2).value == "Answer 1"
    assert out_wb["CC1"].cell(3, 2).value == "Answer 2"


def test_docx_export_table_placement() -> None:
    """DOCX export places answers in table cells."""
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_paragraph("Header")
    t = doc.add_table(rows=3, cols=2)
    t.rows[0].cells[0].text = "Question"
    t.rows[0].cells[1].text = "Answer"
    t.rows[1].cells[0].text = "Q1?"
    t.rows[1].cells[1].text = ""
    t.rows[2].cells[0].text = "Q2?"
    t.rows[2].cells[1].text = ""
    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    questions = [
        {"id": 1, "source_location": json.dumps({"type": "table", "table": 1, "row": 2, "col": 1})},
        {"id": 2, "source_location": json.dumps({"type": "table", "table": 1, "row": 3, "col": 1})},
    ]
    qa = {1: "A1", 2: "A2"}
    out = create_export_from_questionnaire_docx(docx_bytes, questions, qa)
    from app.services.docx_writer import load_document_from_bytes

    doc2 = load_document_from_bytes(out)
    assert doc2.tables[0].rows[1].cells[1].text == "A1"
    assert doc2.tables[0].rows[2].cells[1].text == "A2"


def test_xlsx_format_preserved() -> None:
    """Export preserves cell formatting (EXP-04)."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    c = ws.cell(1, 1, "Q?")
    c.font = openpyxl.styles.Font(bold=True)
    ws.cell(1, 2, "")
    buf = BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()

    questions = [{"id": 1, "source_location": json.dumps({"sheet": "Sheet1", "row": 1, "col": 1})}]
    qa = {1: "Yes"}
    out = create_export_from_questionnaire(xlsx_bytes, questions, qa)
    out_wb = openpyxl.load_workbook(BytesIO(out), data_only=False)
    answer_cell = out_wb["Sheet1"].cell(1, 2)
    assert answer_cell.value == "Yes"
    assert answer_cell.font.bold is True
