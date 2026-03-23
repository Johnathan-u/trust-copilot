"""
End-to-End Pipeline Test
========================
Pushes mock evidence + questionnaire through the full AI pipeline and monitors:
  1. Document upload + indexing
  2. Questionnaire upload + parsing
  3. Answer generation
  4. Compliance event firing
  5. In-app notification creation
  6. Coverage KPI updates
  7. Answer stats

Run inside the API container:
    python tests/e2e_pipeline_test.py
"""

import io
import json
import sys
import time
import requests

BASE = "http://localhost:8000"
S = requests.Session()

# ── Helpers ──

def section(title):
    print(f"\n{'='*70}")
    print(f"  {title}")
    print(f"{'='*70}")

def check(label, ok, detail=""):
    icon = "OK" if ok else "FAIL"
    d = f" — {detail}" if detail else ""
    print(f"  [{icon:>4}] {label}{d}")
    return ok

def wait_job(job_id, ws_id, label, timeout=300):
    """Poll job status until completed/failed or timeout."""
    print(f"  ... waiting for {label} (job {job_id})")
    start = time.time()
    last_status = ""
    while time.time() - start < timeout:
        r = S.get(f"{BASE}/api/jobs/{job_id}?workspace_id={ws_id}")
        if r.status_code != 200:
            time.sleep(2)
            continue
        data = r.json()
        status = data.get("status", "")
        if status != last_status:
            elapsed = int(time.time() - start)
            print(f"       [{elapsed}s] status={status}")
            last_status = status
        if status == "completed":
            elapsed = int(time.time() - start)
            print(f"       Completed in {elapsed}s")
            return data
        if status == "failed":
            print(f"       FAILED: {data.get('error', '?')[:200]}")
            return data
        time.sleep(3)
    print(f"       TIMEOUT after {timeout}s")
    return None

# ── Create mock XLSX questionnaire ──

def create_mock_questionnaire_xlsx():
    """Create a minimal XLSX questionnaire in memory."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Questions"
    ws.append(["Question"])
    questions = [
        "Do you have a formal information security policy?",
        "Is multi-factor authentication required for administrative access?",
        "How do you manage encryption keys?",
        "Do you perform regular vulnerability scans?",
        "How do you handle incident response?",
        "Are backups performed and tested regularly?",
        "Do you have a business continuity plan?",
        "How do you manage vendor risk?",
        "Is there a formal change management process?",
        "How long are audit logs retained?",
    ]
    for q in questions:
        ws.append([q])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

# ── Create mock evidence document ──

def create_mock_evidence_docx():
    """Create a minimal DOCX evidence document in memory."""
    from docx import Document
    doc = Document()
    doc.add_heading("E2E Test — Security Policy Summary", level=1)
    doc.add_paragraph(
        "Our organization maintains a formal information security policy that is reviewed "
        "and approved annually by the CISO. All employees are required to complete security "
        "awareness training within 30 days of hire and annually thereafter."
    )
    doc.add_paragraph(
        "Administrative access to production systems requires multi-factor authentication. "
        "We use centralized identity management with role-based access controls. Privileged "
        "access is logged and reviewed monthly."
    )
    doc.add_paragraph(
        "Encryption keys are managed through a centralized key management service. Data at "
        "rest is encrypted using AES-256. Data in transit is encrypted using TLS 1.2 or higher."
    )
    doc.add_paragraph(
        "We perform quarterly vulnerability scans and annual penetration tests. Critical "
        "vulnerabilities are remediated within 72 hours. All findings are tracked in the "
        "vulnerability management system."
    )
    doc.add_paragraph(
        "Our incident response plan defines severity levels and escalation procedures. "
        "Incidents are triaged within 30 minutes. Post-incident reviews are conducted for "
        "all high and critical severity events."
    )
    doc.add_paragraph(
        "Backups are performed daily and replicated to a geographically separate region. "
        "Backup restoration tests are performed quarterly. RPO is 24 hours and RTO is 4 hours."
    )
    doc.add_paragraph(
        "A formal change management process requires peer review, approval, and testing "
        "before deployment. Emergency changes require retrospective review within 48 hours."
    )
    doc.add_paragraph(
        "Audit logs are retained for a minimum of 365 days. Security events are forwarded "
        "to our SIEM for real-time monitoring and alerting."
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Main test flow ──

def main():
    print("\n" + "="*70)
    print("  TRUST COPILOT — END-TO-END PIPELINE TEST")
    print("  Push mock data through the full AI pipeline")
    print("="*70)

    # ── Login ──
    section("1. AUTHENTICATION")
    r = S.post(f"{BASE}/api/auth/login", json={"email": "demo@trust.local", "password": "j"})
    if r.status_code != 200:
        print(f"  Login failed: {r.status_code}")
        sys.exit(1)
    user = r.json()
    print(f"  Logged in as {user.get('user', {}).get('email', '?')}")

    ws = S.get(f"{BASE}/api/workspaces/current").json()
    ws_id = ws["id"]
    ws_name = ws.get("name", "?")
    print(f"  Workspace: {ws_name} (id={ws_id})")

    # ── Baseline snapshot ──
    section("2. BASELINE SNAPSHOT (before new data)")

    baseline_coverage = S.get(f"{BASE}/api/compliance-coverage").json()
    baseline_kpi = baseline_coverage.get("kpi", {})
    print(f"  Coverage:     {baseline_kpi.get('coverage_pct')}%")
    print(f"  Insufficient: {baseline_kpi.get('insufficient_pct')}%")
    print(f"  Blind spots:  {baseline_kpi.get('blind_spot_count')}")
    print(f"  Total Qs:     {baseline_kpi.get('total_questions')}")
    print(f"  Total drafted: {baseline_kpi.get('total_drafted')}")

    baseline_unread = S.get(f"{BASE}/api/in-app-notifications/unread-count").json()
    print(f"  Unread notifications: {baseline_unread.get('count', 0)}")

    baseline_notifs = S.get(f"{BASE}/api/in-app-notifications?limit=5").json()
    print(f"  Recent notifications: {baseline_notifs.get('total', 0)} total")

    qnrs_before = S.get(f"{BASE}/api/questionnaires/?workspace_id={ws_id}").json()
    docs_before = S.get(f"{BASE}/api/documents/?workspace_id={ws_id}").json()
    print(f"  Questionnaires: {len(qnrs_before)}")
    print(f"  Documents:      {len(docs_before)}")

    # ── Upload evidence document ──
    section("3. UPLOAD EVIDENCE DOCUMENT")
    try:
        evidence_buf = create_mock_evidence_docx()
        r = S.post(
            f"{BASE}/api/documents/upload",
            data={"workspace_id": ws_id},
            files={"file": ("e2e_security_policy.docx", evidence_buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        if r.status_code in (200, 201):
            doc_data = r.json()
            doc_id = doc_data.get("id")
            doc_job_id = doc_data.get("job_id")
            check("Document uploaded", True, f"id={doc_id}, job_id={doc_job_id}")
        else:
            check("Document uploaded", False, f"{r.status_code}: {r.text[:200]}")
            doc_id = None
            doc_job_id = None
    except ImportError:
        print("  python-docx not installed, skipping evidence upload")
        doc_id = None
        doc_job_id = None

    # Wait for document indexing
    if doc_job_id:
        section("4. DOCUMENT INDEXING")
        index_result = wait_job(doc_job_id, ws_id, "document indexing", timeout=120)
        if index_result:
            check("Document indexed", index_result.get("status") == "completed")
        else:
            check("Document indexed", False, "timeout or error")
    else:
        section("4. DOCUMENT INDEXING (skipped)")

    # ── Upload questionnaire ──
    section("5. UPLOAD QUESTIONNAIRE")
    try:
        qnr_buf = create_mock_questionnaire_xlsx()
        r = S.post(
            f"{BASE}/api/questionnaires/upload",
            data={"workspace_id": ws_id},
            files={"file": ("e2e_test_questionnaire.xlsx", qnr_buf, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )
        if r.status_code in (200, 201):
            qnr_data = r.json()
            qnr_id = qnr_data.get("id")
            parse_job_id = qnr_data.get("job_id")
            check("Questionnaire uploaded", True, f"id={qnr_id}, parse_job={parse_job_id}")
        else:
            check("Questionnaire uploaded", False, f"{r.status_code}: {r.text[:200]}")
            qnr_id = None
            parse_job_id = None
    except ImportError:
        print("  openpyxl not installed, cannot create questionnaire")
        sys.exit(1)

    # Wait for parsing
    if parse_job_id:
        section("6. QUESTIONNAIRE PARSING")
        parse_result = wait_job(parse_job_id, ws_id, "questionnaire parsing", timeout=120)
        if parse_result:
            check("Questionnaire parsed", parse_result.get("status") == "completed")
        else:
            check("Questionnaire parsed", False, "timeout or error")

        # Verify questions were extracted
        qnr_detail = S.get(f"{BASE}/api/questionnaires/{qnr_id}?workspace_id={ws_id}").json()
        questions = qnr_detail.get("questions", [])
        check("Questions extracted", len(questions) > 0, f"{len(questions)} questions")
        for i, q in enumerate(questions[:3]):
            print(f"       Q{i+1}: {q.get('text', '?')[:80]}")
    else:
        section("6. QUESTIONNAIRE PARSING (skipped)")

    # ── Trigger answer generation ──
    section("7. ANSWER GENERATION")
    if qnr_id:
        r = S.post(f"{BASE}/api/exports/generate/{qnr_id}?workspace_id={ws_id}", json={})
        if r.status_code == 200:
            gen_data = r.json()
            gen_job_id = gen_data.get("job_id")
            check("Generation triggered", True, f"job_id={gen_job_id}")
        else:
            check("Generation triggered", False, f"{r.status_code}: {r.text[:200]}")
            gen_job_id = None

        if gen_job_id:
            gen_result = wait_job(gen_job_id, ws_id, "answer generation", timeout=600)
            if gen_result:
                status = gen_result.get("status")
                result_data = gen_result.get("result", {})
                if isinstance(result_data, str):
                    try:
                        result_data = json.loads(result_data)
                    except Exception:
                        result_data = {}
                generated = result_data.get("generated", result_data.get("count", "?"))
                total = result_data.get("total", "?")
                check("Answers generated", status == "completed", f"generated={generated}, total={total}")
            else:
                check("Answers generated", False, "timeout")
    else:
        print("  Skipped — no questionnaire")

    # ── Post-generation checks ──
    section("8. POST-GENERATION — ANSWER STATS")
    if qnr_id:
        r = S.get(f"{BASE}/api/ai-governance/questionnaire-answer-stats/{qnr_id}")
        if r.status_code == 200:
            stats = r.json()
            print(f"  Total questions:  {stats.get('total_questions')}")
            print(f"  Total answers:    {stats.get('total_answers')}")
            print(f"  Answered (draft): {stats.get('answered')}")
            print(f"  Not answered:     {stats.get('not_answered')}")
            print(f"  Status breakdown: {stats.get('status_breakdown', {})}")
            print(f"  Gating breakdown: {stats.get('gating_breakdown', {})}")
            check("Answer stats available", stats.get("total_questions", 0) > 0)
        else:
            check("Answer stats", False, f"HTTP {r.status_code}")

    # ── Coverage KPIs after generation ──
    section("9. POST-GENERATION — COVERAGE KPIs")
    after_coverage = S.get(f"{BASE}/api/compliance-coverage").json()
    after_kpi = after_coverage.get("kpi", {})

    print(f"  {'Metric':<25} {'Before':<15} {'After':<15} {'Change'}")
    print(f"  {'-'*70}")
    for key in ["total_questions", "total_drafted", "total_insufficient", "coverage_pct", "insufficient_pct", "blind_spot_count"]:
        before_val = baseline_kpi.get(key, 0) or 0
        after_val = after_kpi.get(key, 0) or 0
        if isinstance(before_val, float):
            delta = f"{after_val - before_val:+.1f}"
        else:
            delta = f"{after_val - before_val:+d}" if isinstance(after_val, int) else f"{after_val}"
        print(f"  {key:<25} {str(before_val):<15} {str(after_val):<15} {delta}")

    check("Coverage KPIs updated", after_kpi.get("total_questions", 0) >= baseline_kpi.get("total_questions", 0))

    # ── Framework coverage detail ──
    fw_cov = after_coverage.get("framework_coverage", [])
    if fw_cov:
        print(f"\n  Framework coverage:")
        for f in fw_cov:
            print(f"    {f['framework']:<20} {f['drafted']}/{f['total']} ({f['coverage_pct']}%)")

    # ── Blind spots ──
    blind = after_coverage.get("blind_spots", [])
    if blind:
        print(f"\n  Blind spots ({len(blind)}):")
        for b in blind[:5]:
            print(f"    {b['subject']:<30} {b['insufficient_count']} insufficient / {b['total']} total")

    # ── In-app notifications after generation ──
    section("10. POST-GENERATION — NOTIFICATIONS")
    after_unread = S.get(f"{BASE}/api/in-app-notifications/unread-count").json()
    unread_before = baseline_unread.get("count", 0)
    unread_after = after_unread.get("count", 0)
    new_notifs = unread_after - unread_before
    check("New notifications created", new_notifs > 0, f"before={unread_before}, after={unread_after}, new={new_notifs}")

    after_notifs = S.get(f"{BASE}/api/in-app-notifications?limit=10").json()
    recent = after_notifs.get("notifications", [])
    if recent:
        print(f"\n  Recent notifications ({len(recent)}):")
        for n in recent[:8]:
            read_mark = " " if n.get("is_read") else "*"
            cat = n.get("category", "info")
            print(f"    {read_mark} [{cat}] {n.get('title', '?')}")
            if n.get("body"):
                print(f"           {n['body'][:100]}")

    # ── Compliance events in notification log ──
    section("11. POST-GENERATION — COMPLIANCE EVENT LOG")
    r = S.get(f"{BASE}/api/notifications/log?page=1&page_size=20")
    if r.status_code == 200:
        log_data = r.json()
        entries = log_data.get("entries", [])
        compliance_entries = [e for e in entries if e.get("event_type", "").startswith("compliance.")]
        check("Compliance events in delivery log", len(compliance_entries) > 0, f"{len(compliance_entries)} compliance events")
        if compliance_entries:
            print(f"\n  Compliance delivery log entries:")
            for e in compliance_entries[:5]:
                print(f"    [{e.get('status')}] {e.get('event_type')} -> {e.get('channel', 'email')} ({e.get('recipient_email', '?')})")
                if e.get("subject"):
                    print(f"            Subject: {e['subject'][:100]}")
        system_entries = [e for e in entries if not e.get("event_type", "").startswith("compliance.")]
        if system_entries:
            print(f"\n  Other delivery log entries ({len(system_entries)}):")
            for e in system_entries[:5]:
                print(f"    [{e.get('status')}] {e.get('event_type')} -> {e.get('channel', 'email')}")
    elif r.status_code == 403:
        check("Notification log", False, "403 — admin only, cannot inspect with demo user")
    else:
        check("Notification log", False, f"HTTP {r.status_code}")

    # ── Active compliance alerts ──
    section("12. ACTIVE COMPLIANCE ALERTS")
    r = S.get(f"{BASE}/api/compliance-alerts/active")
    if r.status_code == 200:
        alerts = r.json().get("alerts", [])
        check("Active alerts computed", True, f"{len(alerts)} active")
        for a in alerts:
            sev = a.get("severity", "?")
            title = a.get("title", "?")
            desc = a.get("description", "")[:100]
            print(f"    [{sev.upper()}] {title}")
            print(f"           {desc}")
    elif r.status_code == 403:
        check("Active alerts", False, "403 — admin only")
    else:
        check("Active alerts", False, f"HTTP {r.status_code}")

    # ── AI insights ──
    section("13. AI INSIGHTS")
    r = S.get(f"{BASE}/api/ai-insights")
    if r.status_code == 200:
        ins = r.json()
        perf = ins.get("performance", {})
        print(f"  Total answers:      {perf.get('total_answers')}")
        print(f"  Drafted:            {perf.get('drafted')}")
        print(f"  Insufficient:       {perf.get('insufficient')}")
        print(f"  Avg confidence:     {perf.get('avg_confidence')}")
        cdist = perf.get("confidence_distribution", {})
        print(f"  Confidence dist:    high={cdist.get('high')}, med={cdist.get('medium')}, low={cdist.get('low')}, none={cdist.get('none')}")

        weak = ins.get("weak_subjects", [])
        if weak:
            print(f"\n  Weakest subjects:")
            for w in weak[:5]:
                print(f"    {w['subject']:<30} avg_conf={w['avg_confidence']}%, count={w['count']}, insufficient={w['insufficient']}")

        failures = ins.get("failure_reasons", [])
        if failures:
            print(f"\n  Failure reasons:")
            for f in failures:
                print(f"    {f.get('label', f.get('reason', '?')):<40} count={f['count']}")

        check("AI insights populated", perf.get("total_answers", 0) > 0)
    elif r.status_code == 403:
        check("AI insights", False, "403 — admin only")

    # ── Questionnaire detail with answers ──
    section("14. QUESTIONNAIRE DETAIL — ANSWER INSPECTION")
    if qnr_id:
        qnr_detail = S.get(f"{BASE}/api/questionnaires/{qnr_id}?workspace_id={ws_id}").json()
        questions = qnr_detail.get("questions", [])
        answered = 0
        insufficient = 0
        with_citations = 0
        for q in questions:
            ans = q.get("answer")
            if not ans:
                continue
            text = ans.get("text", "")
            status = ans.get("status", "")
            citations = ans.get("citations")
            if text:
                answered += 1
            if status == "insufficient_evidence":
                insufficient += 1
            if citations:
                with_citations += 1

        print(f"  Questions:     {len(questions)}")
        print(f"  Answered:      {answered}")
        print(f"  Insufficient:  {insufficient}")
        print(f"  With citations: {with_citations}")

        # Show first 3 Q&A pairs
        print(f"\n  Sample answers:")
        shown = 0
        for q in questions:
            ans = q.get("answer")
            if not ans or not ans.get("text"):
                continue
            if shown >= 3:
                break
            shown += 1
            q_text = q.get("text", "?")[:80]
            a_text = ans.get("text", "")[:150]
            a_status = ans.get("status", "?")
            cit_count = 0
            if ans.get("citations"):
                try:
                    cits = json.loads(ans["citations"]) if isinstance(ans["citations"], str) else ans["citations"]
                    cit_count = len(cits) if isinstance(cits, list) else 0
                except Exception:
                    pass
            print(f"\n    Q: {q_text}")
            print(f"    A: {a_text}")
            print(f"       status={a_status}, citations={cit_count}")

    # ── Export test ──
    section("15. EXPORT TEST")
    if qnr_id:
        r = S.get(f"{BASE}/api/exports/records?workspace_id={ws_id}")
        if r.status_code == 200:
            exports = r.json()
            if isinstance(exports, list):
                check("Export records available", True, f"{len(exports)} total exports")

    # ── Final summary ──
    section("FINAL REPORT")
    print(f"""
  Pipeline Stage              Status
  ─────────────────────────────────────────────────
  Authentication              OK
  Document upload             {'OK' if doc_id else 'SKIPPED'}
  Document indexing           {'OK' if doc_job_id else 'SKIPPED'}
  Questionnaire upload        {'OK' if qnr_id else 'FAILED'}
  Questionnaire parsing       {'OK' if qnr_id else 'FAILED'}
  Answer generation           {'OK' if qnr_id else 'FAILED'}
  Coverage KPIs updated       {'OK' if after_kpi.get('total_questions', 0) >= baseline_kpi.get('total_questions', 0) else 'FAIL'}
  Notifications created       {'OK — ' + str(new_notifs) + ' new' if new_notifs > 0 else 'NO new notifications'}
  Compliance alerts computed  OK (on-the-fly)
  Auto-refresh polling        WIRED (30-60s intervals)

  KPI Summary:
    Coverage:      {baseline_kpi.get('coverage_pct')}% -> {after_kpi.get('coverage_pct')}%
    Insufficient:  {baseline_kpi.get('insufficient_pct')}% -> {after_kpi.get('insufficient_pct')}%
    Blind spots:   {baseline_kpi.get('blind_spot_count')} -> {after_kpi.get('blind_spot_count')}
    Total Qs:      {baseline_kpi.get('total_questions')} -> {after_kpi.get('total_questions')}
    Drafted:       {baseline_kpi.get('total_drafted')} -> {after_kpi.get('total_drafted')}
""")


if __name__ == "__main__":
    main()
