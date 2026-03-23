"""Evidence parsers: PDF (DOC-03), DOCX (DOC-04), XLSX (DOC-05)."""

from pathlib import Path
from typing import Iterator


def parse_pdf(path: str | Path) -> Iterator[dict]:
    """Extract text and structure from PDF (DOC-03)."""
    try:
        import pymupdf
    except ImportError:
        return
    doc = pymupdf.open(path)
    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            yield {"page": i + 1, "text": text.strip()}
    doc.close()


def parse_docx(path: str | Path) -> Iterator[dict]:
    """Extract text from DOCX (DOC-04)."""
    try:
        from docx import Document
    except ImportError:
        return
    doc = Document(path)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            yield {"section": i, "text": para.text.strip()}
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                yield {"text": " | ".join(cells)}


def parse_xlsx_evidence(path: str | Path) -> Iterator[dict]:
    """Extract text from XLSX evidence (DOC-05)."""
    from app.services.xlsx_parser import iter_sheets
    for sheet_name, rows in iter_sheets(path):
        for ri, row in enumerate(rows):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                yield {"sheet": sheet_name, "row": ri + 1, "text": " | ".join(cells)}


def parse_plain_text(path: str | Path) -> list[dict]:
    """Extract text from plain .txt for evidence indexing."""
    p = Path(path)
    if not p.exists():
        return []
    text = p.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        return []
    return [{"text": text}]


def parse_evidence(path: str | Path) -> list[dict]:
    """Dispatch to appropriate parser by extension."""
    p = Path(path)
    ext = p.suffix.lower()
    results = []
    if ext == ".pdf":
        results = list(parse_pdf(path))
    elif ext in (".docx", ".doc"):
        results = list(parse_docx(path))
    elif ext in (".xlsx", ".xls"):
        results = list(parse_xlsx_evidence(path))
    elif ext == ".txt":
        results = parse_plain_text(path)
    return results
