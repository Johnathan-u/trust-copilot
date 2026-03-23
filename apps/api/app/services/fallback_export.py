"""Fallback export pack when in-file placement is not possible (EXP-08)."""

import csv
from io import BytesIO
from typing import Any


def create_fallback_pack_docx(
    questions: list[dict],
    question_to_answer: dict[int, str],
) -> bytes:
    """Create a simple DOCX with Question | Answer table."""
    try:
        from docx import Document
    except ImportError:
        return create_fallback_pack_csv(questions, question_to_answer)

    doc = Document()
    doc.add_heading("Questionnaire Answers (Fallback Export)", 0)
    doc.add_paragraph(
        "Direct placement into the source file was not possible. "
        "Below are all questions with their answers."
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Question"
    hdr[1].text = "Answer"
    for q in questions:
        qid = q.get("id")
        if qid is None:
            continue
        text = q.get("text", "")
        answer = question_to_answer.get(qid, "")
        row = table.add_row().cells
        row[0].text = str(text)
        row[1].text = str(answer)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_fallback_pack_csv(
    questions: list[dict],
    question_to_answer: dict[int, str],
) -> bytes:
    """Create CSV with Question, Answer columns."""
    buf = BytesIO()
    writer = csv.writer(buf)
    writer.writerow(["Question", "Answer"])
    for q in questions:
        qid = q.get("id")
        if qid is None:
            continue
        text = q.get("text", "")
        answer = question_to_answer.get(qid, "")
        writer.writerow([text, answer])
    return buf.getvalue()
