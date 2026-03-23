"""DOCX questionnaire parser (QNR-07, QNR-08, QNR-09)."""

import json
from pathlib import Path

from app.services.answer_type import infer_answer_type
from app.services.question_detector import looks_like_question, looks_like_section_header
from app.services.question_normalizer import normalize_question as normalize


def parse_docx_questionnaire(path: str | Path) -> list[dict]:
    """Parse DOCX into questions from paragraphs and tables. Source: table/row/cell or para index (RES-01)."""
    try:
        from docx import Document
    except ImportError:
        return []
    try:
        doc = Document(path)
    except Exception:
        return []
    results = []
    section = None

    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue
        if looks_like_section_header(text):
            section = text
            continue
        if looks_like_question(text):
            norm = normalize(text)
            if len(norm) < 10:
                continue
            source = {"file": str(path), "type": "paragraph", "index": i + 1}
            results.append({
                "text": norm,
                "section": section,
                "answer_type": infer_answer_type(text),
                "source_location": json.dumps(source),
                "confidence": 80,
            })

    for ti, tbl in enumerate(doc.tables):
        tbl_section = section
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                text = (cell.text or "").strip()
                if not text:
                    continue
                if looks_like_section_header(text):
                    tbl_section = text
                    break
                if looks_like_question(text):
                    norm = normalize(text)
                    if len(norm) < 10:
                        continue
                    source = {"file": str(path), "type": "table", "table": ti + 1, "row": ri + 1, "col": ci + 1}
                    results.append({
                        "text": norm,
                        "section": tbl_section,
                        "answer_type": infer_answer_type(text),
                        "source_location": json.dumps(source),
                        "confidence": 80,
                    })
                    break

    return results
