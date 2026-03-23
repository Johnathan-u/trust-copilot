"""DOCX document writer for Word-based exports (EXP-05)."""

import json
from io import BytesIO
from typing import Any


def load_document_from_bytes(data: bytes):
    """Load DOCX from bytes. Raises on corrupt/invalid file (RES-02)."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX export")
    if not data:
        raise ValueError("Empty document data")
    return Document(BytesIO(data))


def save_document_to_bytes(doc) -> bytes:
    """Save DOCX to bytes."""
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def placements_from_questions_and_answers_docx(
    questions: list[dict],
    question_to_answer: dict[int, str],
) -> list[dict[str, Any]]:
    """Build DOCX placement list from questions with source_location.

    source_location formats:
    - table: {type: "table", table: 1-based, row: 1-based, col: 1-based}
    - paragraph: {type: "paragraph", index: 1-based}
    """
    placements = []
    for q in questions:
        qid = q.get("id")
        if qid is None:
            continue
        text = question_to_answer.get(qid, "")
        loc_raw = q.get("source_location")
        if not loc_raw:
            continue
        try:
            loc = json.loads(loc_raw) if isinstance(loc_raw, str) else loc_raw
        except (json.JSONDecodeError, TypeError):
            continue
        if not isinstance(loc, dict):
            continue
        if loc.get("type") == "table":
            placements.append({
                "type": "table",
                "table": loc.get("table", 1),
                "row": loc.get("row", 1),
                "col": loc.get("col", 1),
                "text": text,
            })
        elif loc.get("type") == "paragraph":
            placements.append({
                "type": "paragraph",
                "index": loc.get("index", 1),
                "text": text,
            })
    return placements


def place_answers_into_document(
    doc,
    placements: list[dict[str, Any]],
    answer_col_offset: int = 1,
) -> None:
    """Place answers into DOCX by table cell or paragraph.

    Table: answer goes in (table, row, col + answer_col_offset).
    Paragraph: answer goes in the next paragraph (index + 1, 1-based).
    """
    for p in placements:
        text = p.get("text", "")
        if p.get("type") == "table":
            ti = (p.get("table", 1) or 1) - 1
            ri = (p.get("row", 1) or 1) - 1
            ci = (p.get("col", 1) or 1) - 1
            answer_col = ci + answer_col_offset
            if ti < 0 or ti >= len(doc.tables):
                continue
            tbl = doc.tables[ti]
            if ri < 0 or ri >= len(tbl.rows):
                continue
            row = tbl.rows[ri]
            if answer_col < len(row.cells):
                cell = row.cells[answer_col]
                cell.text = str(text)
        elif p.get("type") == "paragraph":
            # Question at 1-based index; answer goes in next paragraph (0-based index = idx)
            idx = (p.get("index", 1) or 1)
            answer_idx_0based = idx
            if 0 <= answer_idx_0based < len(doc.paragraphs):
                para = doc.paragraphs[answer_idx_0based]
                para.clear()
                para.add_run(str(text))


def create_export_from_questionnaire_docx(
    docx_bytes: bytes,
    questions: list[dict],
    question_to_answer: dict[int, str],
    answer_col_offset: int = 1,
) -> bytes:
    """Load original DOCX, place answers, return modified DOCX bytes."""
    doc = load_document_from_bytes(docx_bytes)
    placements = placements_from_questions_and_answers_docx(questions, question_to_answer)
    place_answers_into_document(doc, placements, answer_col_offset)
    return save_document_to_bytes(doc)
